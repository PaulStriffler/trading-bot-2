"""Generate Trading-Strategie.docx — single A4 page, print-ready.

Layout: A4 portrait, ~1.4 cm margins, 2-column body via a borderless table.
Font: Helvetica/Arial, tight spacing, dark blue accents.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x26, 0x62, 0xFF)        # blue
ACCENT_DARK = RGBColor(0x0D, 0x47, 0xA1)
WIN = RGBColor(0x26, 0xA6, 0x9A)
TEXT_DIM = RGBColor(0x55, 0x55, 0x55)
BORDER_GRAY = "DDDDDD"
HIGHLIGHT_BG = "F1F5FB"


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tc_pr.append(mar)


def tight(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.12


def add_heading(cell, text, size=10.5, color=ACCENT_DARK):
    p = cell.add_paragraph()
    tight(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Helvetica Neue"
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_para(cell, text, size=8.5, bold=False, color=None, indent=False):
    p = cell.add_paragraph()
    tight(p)
    if indent:
        p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.name = "Helvetica Neue"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return p


def add_bullet(cell, text, size=8.5):
    p = cell.add_paragraph()
    tight(p)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run("•  " + text)
    run.font.name = "Helvetica Neue"
    run.font.size = Pt(size)
    return p


def add_numbered_step(cell, num, text, size=8.5):
    p = cell.add_paragraph()
    tight(p)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run_num = p.add_run(f"{num}  ")
    run_num.font.name = "Helvetica Neue"
    run_num.font.size = Pt(size)
    run_num.font.bold = True
    run_num.font.color.rgb = ACCENT
    run_txt = p.add_run(text)
    run_txt.font.name = "Helvetica Neue"
    run_txt.font.size = Pt(size)
    return p


def build():
    doc = Document()

    # ---- Page ----
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(1.3)
    s.bottom_margin = Cm(1.3)
    s.left_margin = Cm(1.5)
    s.right_margin = Cm(1.5)

    # ---- Default style ----
    style = doc.styles["Normal"]
    style.font.name = "Helvetica Neue"
    style.font.size = Pt(8.5)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(title)
    title.paragraph_format.space_after = Pt(1)
    r = title.add_run("Trading-Strategie · ICT/SMC Multi-Timeframe (FINAL v3)")
    r.font.name = "Helvetica Neue"
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = ACCENT_DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(sub)
    sub.paragraph_format.space_after = Pt(6)
    r = sub.add_run("FTMO Funded Swing $10.000  ·  1 % Risk  ·  Hebel 1:30  ·  US500 · US30 · US100 · US2000")
    r.font.name = "Helvetica Neue"
    r.font.size = Pt(9)
    r.font.color.rgb = TEXT_DIM

    # ---- Account-Box (full-width banner) ----
    banner = doc.add_table(rows=1, cols=4)
    banner.autofit = False
    widths = [Cm(4.5), Cm(4.5), Cm(4.5), Cm(4.5)]
    for i, w in enumerate(widths):
        banner.columns[i].width = w
    for col_idx, (label, val) in enumerate([
        ("Account", "$10.000"),
        ("Risk / Trade", "1 % = $100"),
        ("Hebel", "1:30"),
        ("Partials TP", "33 / 33 / 34"),
    ]):
        cell = banner.cell(0, col_idx)
        cell.width = widths[col_idx]
        set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
        cell_shading(cell, HIGHLIGHT_BG)
        for side in ("top", "left", "bottom", "right"):
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders")) or OxmlElement("w:tcBorders")
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:color"), BORDER_GRAY)
            borders.append(b)
            if borders.getparent() is None:
                tc_pr.append(borders)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tight(p1)
        r = p1.add_run(label.upper())
        r.font.name = "Helvetica Neue"; r.font.size = Pt(7); r.font.color.rgb = TEXT_DIM; r.font.bold = True
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tight(p2)
        r = p2.add_run(val)
        r.font.name = "Helvetica Neue"; r.font.size = Pt(10.5); r.font.bold = True

    spacer = doc.add_paragraph(); tight(spacer); spacer.paragraph_format.space_after = Pt(4)

    # ---- Two-column body table ----
    body = doc.add_table(rows=1, cols=2)
    body.autofit = False
    body.columns[0].width = Cm(9.0)
    body.columns[1].width = Cm(9.0)
    left = body.cell(0, 0); right = body.cell(0, 1)
    left.width = Cm(9.0); right.width = Cm(9.0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    remove_cell_borders(left); remove_cell_borders(right)
    set_cell_margins(left, left=0, right=180)
    set_cell_margins(right, left=180, right=0)

    # Remove the default empty paragraph in each cell so we control spacing
    left.paragraphs[0]._element.getparent().remove(left.paragraphs[0]._element)
    right.paragraphs[0]._element.getparent().remove(right.paragraphs[0]._element)

    # ===== LEFT COLUMN =====
    add_heading(left, "0 · Wochenvorbereitung (Sonntag)")
    add_bullet(left, "Levels markieren: Order Blocks, BPRs (auch inverse), Fair Value Gaps (auch inverse).")
    add_bullet(left, "Inverse FVG: Gap wird ohne Reaktion durchbrochen, Preis kehrt in die Zone zurück und reagiert von dort in die entgegengesetzte Richtung der ursprünglichen FVG.")

    add_heading(left, "1 · Weekly Bias")
    add_bullet(left, "Richtung des letzten Break of Structure auf Weekly")
    add_bullet(left, "ODER: 5 aufeinanderfolgende Closes (bull/bear)")
    add_bullet(left, "Bias bull → nur Long-Setups · Bias bear → nur Short-Setups")

    add_heading(left, "2 · Daily Liquidity Sweep")
    add_bullet(left, "Long-Setup: Pivot-Low wird kurz untertroffen, Close darüber.")
    add_bullet(left, "Short-Setup: Pivot-High wird kurz überschritten, Close darunter.")

    add_heading(left, "3 · Order Block (4H)")
    add_bullet(left, "Letzter gegensätzlich-farbiger Move-Block VOR dem Sweep markieren.")
    add_bullet(left, "Bei mehrfachem Sweep: alle 3 OBs prüfen, denjenigen mit den besten Confluences wählen.")

    # ===== RIGHT COLUMN =====
    add_heading(right, "4 · Entry-Sequenz (1H) — KERN")
    add_numbered_step(right, "①", "Preis muss AUS dem Order Block RAUSKOMMEN.")
    add_numbered_step(right, "②", "Preis muss WIEDER REINKOMMEN in den OB.")
    add_numbered_step(right, "③", "Im OB muss eine Confluence (FVG / inverse FVG / BPR) entstehen — darf den OB überlappen.")
    add_numbered_step(right, "④", "BOS RAUS aus der Confluence (nicht aus dem OB).")
    add_numbered_step(right, "⑤", "ENTRY am Close der BOS-Bestätigung.")

    add_heading(right, "5 · Stop Loss")
    add_bullet(right, "Unter dem OB (Long) / über dem OB (Short).")
    add_bullet(right, "Puffer: 0.5 × ATR(4H).  Beispiel: ATR(4H)=100 → SL = OB-Low − 50.")

    add_heading(right, "6 · Take Profits (3 Stufen)")
    add_bullet(right, "TP1 = 1 : 1 RR  (sichert Break-even).")
    add_bullet(right, "TP3 = nächste Previous Area of Liquidity (PAL) in Gewinnrichtung.")
    add_bullet(right, "TP2 = Mitte zwischen TP1 und TP3.")
    add_bullet(right, "Plausibilität: wenn TP1 ≈ 2 %, dann TP3 ≥ 3 %.")
    add_bullet(right, "Wenn keine PAL: Fib-Extension auf Daily (Swing-Low → Swing-High → Low vor Entry), Ziele 1.272 / 1.618 / 2.0.")

    add_heading(right, "7 · Risk Management")
    add_bullet(right, "1 % Risk pro Trade ($100 bei $10k Account).")
    add_bullet(right, "Partial-Close: 33 / 33 / 34 %.")
    add_bullet(right, "Nach TP1 Treffer: SL auf Break-even.")
    add_bullet(right, "Daily Kill-Switch: −3 % Tagesverlust → Stopp.")

    # ---- Footer (one-line summary) ----
    foot_spacer = doc.add_paragraph(); tight(foot_spacer); foot_spacer.paragraph_format.space_before = Pt(4)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(foot)
    r = foot.add_run("Weekly Bias → Daily Sweep → 4H Order Block → 1H raus/rein/FVG/BOS → Entry → SL/TPs")
    r.font.name = "Helvetica Neue"; r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = TEXT_DIM

    out = Path(__file__).parent / "Trading-Strategie.docx"
    doc.save(out)
    print(f"Wrote {out}  ({out.stat().st_size//1024} KB)")
    return out


if __name__ == "__main__":
    build()
